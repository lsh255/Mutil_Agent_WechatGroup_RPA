from docx import Document
from jinja2 import Template
from typing import Dict, Any
from pathlib import Path


class GenerateReportTool:
    """Word报告生成工具：用于生成工作报告"""
    
    def __init__(self, template_path: str, output_dir: str):
        """初始化报告生成工具
        
        Args:
            template_path: Jinja2模板文件路径
            output_dir: 输出目录
        """
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _load_template(self) -> Template:
        """加载Jinja2模板
        
        Returns:
            Template对象
        """
        if not self.template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {self.template_path}")
        
        with open(self.template_path, "r", encoding="utf-8") as f:
            template_content = f.read()
        
        return Template(template_content)
    
    def generate(self, data: Dict[str, Any], output_filename: str = None) -> str:
        """生成Word报告
        
        Args:
            data: 报告数据字典，应包含:
                - user: 用户名
                - task_id: 任务ID
                - summary: 内容摘要
                - extracted_info: 提取的信息
            output_filename: 输出文件名（可选）
            
        Returns:
            生成的报告文件路径
        """
        try:
            # 加载模板
            template = self._load_template()
            
            # 渲染模板
            report_content = template.render(**data)
            
            # 确定输出文件名
            if not output_filename:
                output_filename = f"report_{data['task_id']}.docx"
            
            output_path = self.output_dir / output_filename
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            doc.add_heading(f"工作报告 - {data['user']}", level=1)
            
            # 添加内容
            doc.add_heading("任务摘要", level=2)
            doc.add_paragraph(data['summary'])
            
            # 添加提取的信息
            if data.get('extracted_info'):
                doc.add_heading("详细信息", level=2)
                for key, value in data['extracted_info'].items():
                    doc.add_paragraph(f"{key}: {value}")
            
            # 保存文档
            doc.save(output_path)
            
            return str(output_path)
            
        except Exception as e:
            print(f"报告生成失败: {e}")
            return ""
    
    def generate_from_template(self, data: Dict[str, Any], output_filename: str = None) -> str:
        """从Word模板生成报告（高级功能）
        
        Args:
            data: 报告数据字典
            output_filename: 输出文件名
            
        Returns:
            生成的报告文件路径
        """
        try:
            # 确定输出文件名
            if not output_filename:
                output_filename = f"report_{data['task_id']}.docx"
            
            output_path = self.output_dir / output_filename
            
            # 从模板创建新文档
            if self.template_path.exists() and self.template_path.suffix == ".docx":
                doc = Document(self.template_path)
            else:
                doc = Document()
            
            # 替换文档中的占位符
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            # 保存文档
            doc.save(output_path)
            
            return str(output_path)
            
        except Exception as e:
            print(f"从模板生成报告失败: {e}")
            return ""
